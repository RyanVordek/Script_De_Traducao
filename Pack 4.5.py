# -*- coding: utf-8 -*-

"""
SUPER TRADUTOR AUTOMÁTICO (v7.3 - Versão Final Consolidada)

Esta versão combina todas as funcionalidades e correções, incluindo:
- Proteção de tags Ren'Py ({...}) e variáveis ([...]).
- Detecção de diálogos com múltiplas expressões.
- Lógica flexível para pareamento de linhas de tradução.
- Tradução de arquivos .docx.
- Adaptação de formalidade e correções contextuais.
"""

import os
import sys
import re
import shutil
import argparse
from functools import lru_cache
import subprocess # Importar o módulo subprocess
import tkinter as tk
from tkinter import filedialog, messagebox, ttk

# --- Bloco de Importação e Verificação de Dependências ---
def install_and_import(package, install_name=None):
    if install_name is None:
        install_name = package
    try:
        __import__(package)
        print(f"✅ Dependência '{package}' já instalada.")
    except ImportError:
        print("=" * 70)
        print(f"❌ A biblioteca '{package}' não foi encontrada.")
        print(f"   Tentando instalar '{install_name}' automaticamente...")
        try:
            # Usar sys.executable para garantir que pip do ambiente correto seja usado
            subprocess.check_call([sys.executable, "-m", "pip", "install", install_name])
            print(f"✅ '{install_name}' instalado com sucesso!")
            __import__(package) # Tentar importar novamente após a instalação
        except subprocess.CalledProcessError as e:
            print(f"❌ ERRO ao instalar '{install_name}': {e}")
            print("   Por favor, tente instalar manualmente:")
            print(f"   pip install {install_name}")
            sys.exit(1)
        except ImportError:
            print(f"❌ ERRO: '{package}' não pôde ser importado mesmo após a instalação.")
            print("   Pode haver um problema com sua instalação Python/pip.")
            sys.exit(1)
        print("=" * 70)

install_and_import('argostranslate')
install_and_import('docx', install_name='python-docx') # python-docx é o nome do pacote para pip

from argostranslate import package, translate
import docx

# --- Variáveis Globais ---
TRADUTOR = None
USAR_FILTRO_GENERO = True
MAPA_GENERO_PERSONAGEM = {}
GLOSSARIO_PROTEGIDO = []
PADROES_PROTEGIDOS = []

PADROES_PROTEGIDOS_PADRAO = [
    r'%\([^)]+\)[sd]',
    r'%[sd]',
    r'\b\d{1,2}:\d{2}\b',
    r'\b\d+(?:[.,]\d+)?\b',
    r'\b[A-Z]{2,}\b',
    r'\b\w+_\w+\b',
    r'\b[\w\-]+\.(?:png|jpg|jpeg|gif|webp|mp3|ogg|wav|webm|mp4)\b',
    r'\b#[0-9A-Fa-f]{3,6}\b',
    r'[A-Za-z]:\\[^\s]+',
    r'/[^\s]+',
]

PADROES_PRONOMES_FEMININOS = [
    (r'\bela\b', 'ele'),
    (r'\bdela\b', 'dele'),
    (r'\bnela\b', 'nele'),
    (r'\bminha\b', 'meu'),
    (r'\bminhas\b', 'meus'),
    (r'\bsua\b', 'seu'),
    (r'\bsuas\b', 'seus'),
    (r'\buma\b', 'um'),
    (r'\buma\s+das\b', 'um dos'),
    (r'\bboa\b', 'bom'),
    (r'\bbonita\b', 'bonito'),
]

RELATORIO_PROBLEMAS = []

# --- Funções Principais de Tradução e Adaptação ---

def configurar_tradutor(codigo_origem="en", codigo_destino="pb", diretorio_base="."):
    """
    Configura o tradutor do Argos Translate, procurando por pacotes locais
    (.argosmodel) antes de usar os pacotes do sistema.
    """
    print(f"[*] Configurando tradutor de '{codigo_origem}' para '{codigo_destino}'...")

    if codigo_destino == "pb":
        # Argos Translate usa 'pt' para Português, que geralmente inclui os modelos para pt-BR.
        codigo_destino = "pb" # Corrigido para 'pt' conforme o Argos Translate usa.

    caminho_pacote_local = None
    try:
        for nome_arquivo in os.listdir(diretorio_base):
            if nome_arquivo.endswith(".argosmodel"):
                caminho_pacote_local = os.path.join(diretorio_base, nome_arquivo)
                print(f"[*] Pacote de idioma local encontrado: {nome_arquivo}")
                break
    except FileNotFoundError:
        print(f"⚠️  Aviso: O diretório base '{diretorio_base}' não foi encontrado para procurar pacotes locais.")


    if caminho_pacote_local:
        try:
            print("[*] Instalando pacote local...")
            package.install_from_path(caminho_pacote_local)
            print("[*] Pacote local instalado com sucesso!")
        except Exception as e:
            print(f"❌ ERRO ao instalar o pacote local: {e}")
            print("    O script tentará usar os pacotes já existentes no sistema.")
    else:
        print("[*] Nenhum pacote de idioma local (.argosmodel) encontrado. Verificando pacotes do sistema.")

    try:
        package.update_package_index()
    except Exception as e:
        print(f"⚠️  Aviso: Não foi possível atualizar a lista de pacotes da internet: {e}")

    idiomas_instalados = translate.get_installed_languages()
    lang_origem = next((lang for lang in idiomas_instalados if lang.code == codigo_origem), None)
    lang_destino = next((lang for lang in idiomas_instalados if lang.code == codigo_destino), None)

    if not lang_origem or not lang_destino:
        print("=" * 70)
        print(f"❌ ERRO CRÍTICO: Pacotes de idioma '{codigo_origem}' -> '{codigo_destino}' não foram carregados.")
        print("   Verifique se há um arquivo .argosmodel válido na pasta ou se os pacotes")
        print("   estão instalados corretamente no sistema (use: argos-translate-gui).")
        # Tentativa de download automático do pacote se não for encontrado e não houver pacote local
        if not caminho_pacote_local:
            print(f"   Tentando baixar o pacote de idioma '{codigo_origem}' para '{codigo_destino}'...")
            try:
                available_packages = package.get_available_packages()
                desired_package = next((p for p in available_packages if p.from_code == codigo_origem and p.to_code == codigo_destino), None)
                if desired_package:
                    package.install_from_path(desired_package.download())
                    print(f"✅ Pacote '{codigo_origem}' para '{codigo_destino}' baixado e instalado com sucesso!")
                    # Recarregar idiomas após a instalação
                    idiomas_instalados = translate.get_installed_languages()
                    lang_origem = next((lang for lang in idiomas_instalados if lang.code == codigo_origem), None)
                    lang_destino = next((lang for lang in idiomas_instalados if lang.code == codigo_destino), None)
                    if lang_origem and lang_destino:
                        print("[*] Tradutor configurado com sucesso após o download.")
                        return lang_origem.get_translation(lang_destino)
                else:
                    print("❌ Não foi possível encontrar o pacote de idioma online.")
            except Exception as download_e:
                print(f"❌ ERRO ao tentar baixar o pacote de idioma: {download_e}")

        sys.exit(1)

    print("[*] Tradutor configurado com sucesso.")
    return lang_origem.get_translation(lang_destino)

@lru_cache(maxsize=16384)
def traduzir_com_cache(texto_original):
    """
    Traduz um texto usando o tradutor global e armazena o resultado em cache.
    Ignora a string 'EMPTYSTRING'.
    """
    if not texto_original or not texto_original.strip() or texto_original == 'EMPTYSTRING':
        return texto_original
    try:
        return TRADUTOR.translate(texto_original)
    except Exception as e:
        print(f"⚠️  Erro ao traduzir o texto '{texto_original[:50]}...': {e}")
        return texto_original

def carregar_glossario(caminho_glossario):
    """
    Carrega uma lista de termos protegidos a partir de um arquivo de texto.
    Linhas em branco e comentários iniciados por # são ignorados.
    """
    termos = []
    if not caminho_glossario:
        return termos
    try:
        with open(caminho_glossario, 'r', encoding='utf-8') as f:
            for linha in f:
                linha = linha.strip()
                if not linha or linha.startswith('#'):
                    continue
                termos.append(linha)
    except FileNotFoundError:
        print(f"⚠️  Glossário não encontrado: {caminho_glossario}")
    return termos

def carregar_padroes_protegidos(caminho_padroes):
    """
    Carrega padrões regex adicionais a partir de um arquivo de texto.
    Linhas em branco e comentários iniciados por # são ignorados.
    """
    padroes = []
    if not caminho_padroes:
        return padroes
    try:
        with open(caminho_padroes, 'r', encoding='utf-8') as f:
            for linha in f:
                linha = linha.strip()
                if not linha or linha.startswith('#'):
                    continue
                padroes.append(linha)
    except FileNotFoundError:
        print(f"⚠️  Padrões de regex não encontrados: {caminho_padroes}")
    return padroes

def carregar_mapa_genero(caminho_genero):
    """
    Carrega um mapa de gênero por personagem no formato:
    personagem=masc|fem|neutro
    """
    mapa = {}
    if not caminho_genero:
        return mapa
    try:
        with open(caminho_genero, 'r', encoding='utf-8') as f:
            for linha in f:
                linha = linha.strip()
                if not linha or linha.startswith('#') or '=' not in linha:
                    continue
                personagem, genero = linha.split('=', 1)
                mapa[personagem.strip()] = genero.strip().lower()
    except FileNotFoundError:
        print(f"⚠️  Mapa de gênero não encontrado: {caminho_genero}")
    return mapa

def proteger_termos(texto):
    """
    Substitui termos do glossário e padrões por placeholders para evitar tradução.
    Retorna o texto protegido e o mapa de restauração.
    """
    substituicoes = {}
    contador = 0

    def substituir_term(match):
        nonlocal contador
        termo = match.group(0)
        chave = f"__TERM_{contador:04d}__"
        substituicoes[chave] = termo
        contador += 1
        return chave

    texto_protegido = texto
    for termo in GLOSSARIO_PROTEGIDO:
        if not termo:
            continue
        padrao = re.compile(r'\b' + re.escape(termo) + r'\b')
        texto_protegido = padrao.sub(substituir_term, texto_protegido)

    for padrao in PADROES_PROTEGIDOS:
        texto_protegido = re.sub(padrao, substituir_term, texto_protegido)

    return texto_protegido, substituicoes

def restaurar_termos(texto, substituicoes):
    for chave, termo in substituicoes.items():
        texto = texto.replace(chave, termo)
    return texto

def aplicar_filtro_genero(texto_traduzido, genero):
    """
    Aplica um filtro para evitar pronomes femininos quando o gênero do personagem
    é masculino ou neutro.
    """
    if genero == 'fem':
        return texto_traduzido
    texto_corrigido = texto_traduzido
    for padrao, substituto in PADROES_PRONOMES_FEMININOS:
        texto_corrigido = re.sub(padrao, substituto, texto_corrigido, flags=re.IGNORECASE)
    return texto_corrigido

def registrar_problemas(texto_original, texto_traduzido, origem):
    """
    Registra ocorrências de possíveis problemas no relatório final.
    """
    problemas = []
    for padrao, _ in PADROES_PRONOMES_FEMININOS:
        if re.search(padrao, texto_traduzido, flags=re.IGNORECASE):
            problemas.append("pronome_feminino")
            break
    if problemas:
        RELATORIO_PROBLEMAS.append({
            "origem": origem,
            "original": texto_original,
            "traduzido": texto_traduzido,
            "problemas": problemas
        })

def traduzir_com_protecao_de_codigo(texto_com_codigo):
    """
    Divide a string em texto e código Ren'Py ({...} ou [...]), traduz apenas o texto
    e depois junta tudo novamente. Essencial para não corromper o jogo.
    """
    # Expressão regular que captura tags com chaves OU variáveis com colchetes.
    padrao_codigo = r'({[^}]+}|\[[^\]]+\])'
    
    # Se não houver nenhum padrão de código, traduz a string inteira.
    if not re.search(padrao_codigo, texto_com_codigo):
        return traduzir_com_cache(texto_com_codigo)

    partes = re.split(padrao_codigo, texto_com_codigo)
    partes_traduzidas = []

    for parte in partes:
        if not parte:
            continue
        # Verifica se a parte é uma tag ou uma variável
        e_tag = parte.startswith('{') and parte.endswith('}')
        e_variavel = parte.startswith('[') and parte.endswith(']')
        
        if e_tag or e_variavel:
            # Se for código, mantém original
            partes_traduzidas.append(parte)
        else:
            # Se for texto, protege termos e traduz
            parte_protegida, substituicoes = proteger_termos(parte)
            traducao = traduzir_com_cache(parte_protegida)
            partes_traduzidas.append(restaurar_termos(traducao, substituicoes))

    return "".join(partes_traduzidas)

def detectar_formalidade_ingles(texto):
    """
    Analisa o texto em inglês usando um sistema de pontuação ponderada para
    determinar o tom com alta precisão.
    """
    PONTUACAO_FORMALIDADE = {
        'sincerely': 5, 'yours faithfully': 5, 'to whom it may concern': 5, 'regards': 4,
        'esteemed': 4, 'mr.': 3, 'mrs.': 3, 'ms.': 3, 'madam': 4, 'sir': 4, 'furthermore': 3,
        'consequently': 3, 'nevertheless': 3, 'henceforth': 3, 'therefore': 2, 'additionally': 2,
        'moreover': 2, 'subsequently': 2, 'thus': 2, 'inquire': 3, 'procure': 3, 'endeavor': 3,
        'commence': 2, 'facilitate': 2, 'ascertain': 2, 'request': 1, 'require': 1, 'assistance': 2,
        'clarification': 2, 'gratitude': 2, 'opportunity': 1, 'documentation': 2, 'pertaining': 2,
        'shall': 2, 'kindly': 1, 'lmao': -5, 'rofl': -5, 'omg': -4, 'btw': -3, 'fyi': -3, 'imo': -3,
        'lol': -3, 'ain\'t': -4, 'cuz': -3, 'gonna': -2, 'wanna': -2, 'gotta': -2, 'dunno': -2,
        'lemme': -2, 'gimme': -2, 'can\'t': -1, 'don\'t': -1, 'won\'t': -1, 'i\'m': -1, 'you\'re': -1,
        'dude': -3, 'bro': -3, 'yo': -3, 'sup': -3, 'what\'s up': -3, 'my bad': -3, 'for real': -2,
        'no worries': -2, 'hang out': -2, 'chill': -2, 'awesome': -2, 'dope': -2, 'sick': -2,
        'lit': -2, 'cool': -1, 'man': -2, 'buddy': -2, 'pal': -2, 'folks': -1, 'hey': -1,
        'yeah': -1, 'yep': -1
    }
    LIMIAR_NEUTRO = 3
    texto_lower = texto.lower()
    pontuacao_total = 0
    for termo, peso in PONTUACAO_FORMALIDADE.items():
        padrao = r'\b' + re.escape(termo) + r'\b'
        ocorrencias = re.findall(padrao, texto_lower)
        pontuacao_total += len(ocorrencias) * peso

    if pontuacao_total > LIMIAR_NEUTRO:
        return 'formal'
    elif pontuacao_total < -LIMIAR_NEUTRO:
        return 'informal'
    else:
        return 'neutro'

def aplicar_correcoes_contextuais_ptbr(texto_traduzido):
    """
    Aplica um conjunto de regras de substituição para corrigir traduções literais
    e outros erros comuns.
    """
    REGRAS_CORRECAO_CONTEXTUAL = {
        r'\b(eu\s+congelo|tô\s+congelando)\b': 'Eu paro',
        r'\bfazem\s+uma\s+abelha\b': 'vão direto',
        r'\b(em\s+sua\s+mãe|na\s+sua\s+mãe)\b': 'na boca dele',
        r'\bcolocando\s+o\s+bolo\s+em\s+sua\s+mãe\b': 'colocando o bolo na boca',
        r'\bgarras\s+copulam\b': 'garras perfuram',
        r'Quieres\s+peloar!\s+Sem\s+mim\s+jodas!': 'Quer brigar! Não fode comigo!',
        r'\bnão\s+na\s+véi\b': 'no rosto não, cara!',
        r'\bpode\s+tocar\s+na\s+minha\s+porta\b': 'Fique à vontade.',
        r'pegar\s+a\s+buzina\s+de\s+alguém': 'pegar o chifre de alguém',
        r'véi\s+de\s+puta': 'filho da puta',
    }
    texto_corrigido = texto_traduzido
    for padrao_errado, correcao in REGRAS_CORRECAO_CONTEXTUAL.items():
        texto_corrigido = re.sub(padrao_errado, correcao, texto_corrigido, flags=re.IGNORECASE)
    return texto_corrigido

def aplicar_adaptacao_ptbr(texto_traduzido, formalidade):
    """
    Adapta o texto ao nível de formalidade detectado (formal ou informal/gírias).
    """
    if formalidade == 'neutro':
        return texto_traduzido

    regras_formais = {
        'você': 'o senhor/a senhora', 'vocês': 'os senhores/as senhoras', 'te': 'lhe',
        'seu': 'seu/sua', 'a gente': 'nós', 'meu': 'meu/minha', 'ajudar': 'auxiliar',
        'precisa': 'necessita', 'conseguir': 'obter', 'pedir': 'solicitar', 'mostrar': 'demonstrar',
        'usar': 'utilizar', 'começar': 'iniciar', 'terminar': 'finalizar', 'dar': 'fornecer',
        'falar': 'comunicar', 'entender': 'compreender', 'ir': 'dirigir-se', 'mandar': 'enviar',
        'querer': 'desejar', 'ver': 'observar', 'dizer': 'declarar', 'achar': 'considerar',
        'confirmar': 'ratificar', 'explicar': 'elucidar', 'morar': 'residir', 'comprar': 'adquirir',
        'pedir desculpas': 'apresentar escusas', 'ajuda': 'auxílio', 'obrigado': 'grato',
        'obrigada': 'grata', 'desculpe': 'lamento', 'coisa': 'questão', 'mas': 'porém',
        'então': 'portanto', 'muito': 'sobremaneira', 'casa': 'residência', 'fim': 'término',
        'conversa': 'diálogo', 'dono': 'proprietário', 'também': 'outrossim',
        'por isso': 'destarte', 'chefe': 'superior'
    }

    regras_informais = {
        'você': 'cê', 'está': 'tá', 'estou': 'tô', 'estamos': 'tamo', 'para': 'pra', 'para o': 'pro',
        'para a': 'pra', 'qual é': 'qualé', 'com o': 'co', 'com a': 'ca', 'dinheiro': 'grana',
        'trabalho': 'trampo', 'trabalhar': 'trampar', 'legal': 'daora', 'bom': 'massa',
        'muito bom': 'show de bola', 'problema': 'B.O.', 'cara': 'véi', 'amigo': 'parça',
        'entende': 'tá ligado', 'entendeu': 'sacou', 'com certeza': 'demorô', 'garota': 'mina',
        'garoto': 'mano', 'rápido': 'ligeiro', 'entendi': 'saquei', 'vamos embora': 'bora',
        'festa': 'rolê', 'combinado': 'fechou', 'confusão': 'treta', 'conversa': 'papo',
        'espera aí': 'péra', 'mesmo': 'mermo', 'tipo': 'tipo assim', 'de boa': 'sussa',
        'ótimo': 'top', 'se talvez': 'se pá', 'complicado': 'tenso', 'não aguento': 'não tanko',
        'pessoa chata': 'cringe'
    }

    regras = regras_formais if formalidade == 'formal' else regras_informais
    def substituir(match):
        palavra_encontrada = match.group(0)
        substituta = regras.get(palavra_encontrada.lower())
        if not substituta: return palavra_encontrada
        if palavra_encontrada.isupper(): return substituta.upper()
        if palavra_encontrada.istitle(): return substituta.capitalize()
        return substituta
    chaves_ordenadas = sorted(regras.keys(), key=len, reverse=True)
    padrao = re.compile(r'\b(' + '|'.join(re.escape(key) for key in chaves_ordenadas) + r')\b', re.IGNORECASE)
    return padrao.sub(substituir, texto_traduzido)

def processar_paragrafo_completo(texto_original):
    """
    Pipeline completo de tradução para um bloco de texto:
    1. Traduz (protegendo código) -> 2. Corrige -> 3. Detecta formalidade -> 4. Adapta.
    """
    if not texto_original or not texto_original.strip():
        return ""
    
    # Usa a nova função com proteção aprimorada
    traducao_base = traduzir_com_protecao_de_codigo(texto_original)
    traducao_corrigida = aplicar_correcoes_contextuais_ptbr(traducao_base)
    formalidade = detectar_formalidade_ingles(texto_original)
    traducao_final = aplicar_adaptacao_ptbr(traducao_corrigida, formalidade)
    
    return traducao_final

# --- MODO DE TRADUÇÃO: REN'PY (.rpy) ---

def processar_arquivo_rpy(caminho_arquivo):
    """Lógica de tradução para um único arquivo .rpy com todas as correções."""
    print(f"\n📄 Processando: {os.path.basename(caminho_arquivo)}")
    try:
        with open(caminho_arquivo, 'r', encoding='utf-8') as f:
            linhas = f.readlines()
    except Exception as e:
        print(f"❌ Erro ao ler o arquivo: {e}")
        return 0

    RE_OLD = re.compile(r'^\s*old\s*"(?P<texto>.+?)"\s*$')
    RE_NEW_VAZIO = re.compile(r'^(?P<indent>\s*new\s*)""\s*$')
    RE_DIALOGO_COMENTADO = re.compile(r'^\s*#\s*(?P<prefixo>.+?)\s+"(?P<texto>.+?)"\s*$')
    RE_DIALOGO_VAZIO = re.compile(r'^(?P<indent>\s*)(?P<prefixo>.+?)\s+""\s*$')
    RE_NARRACAO_COMENTADA = re.compile(r'^\s*#\s*"(?P<texto>.+?)"(?P<resto>.*)$')
    RE_NARRACAO_VAZIA = re.compile(r'^(?P<indent>\s*)""(?P<resto>.*)$')

    novas_linhas = []
    traducoes_feitas = 0
    i = 0
    total_linhas = len(linhas)

    while i < total_linhas:
        linha_atual = linhas[i]
        proxima_linha = linhas[i + 1] if i + 1 < total_linhas else None
        traduziu_bloco = False

        if proxima_linha:
            texto_original = None
            linha_traduzida_formatada = None

            m_old = RE_OLD.match(linha_atual)
            m_new = RE_NEW_VAZIO.match(proxima_linha)
            m_com_diag = RE_DIALOGO_COMENTADO.match(linha_atual)
            m_vazio_diag = RE_DIALOGO_VAZIO.match(proxima_linha)
            m_com_narr = RE_NARRACAO_COMENTADA.match(linha_atual)
            m_vazio_narr = RE_NARRACAO_VAZIA.match(proxima_linha)

            if m_old and m_new:
                texto_original = m_old.group('texto')
                traducao_final = processar_paragrafo_completo(texto_original)
                linha_traduzida_formatada = f'{m_new.group("indent")}"{traducao_final}"\n'

            elif m_com_diag and m_vazio_diag:
                prefixo_comentado = m_com_diag.group('prefixo')
                prefixo_vazio = m_vazio_diag.group('prefixo')
                if prefixo_comentado.split()[0] == prefixo_vazio.split()[0]:
                    texto_original = m_com_diag.group('texto')
                    personagem = prefixo_vazio.split()[0]
                    traducao_final = processar_paragrafo_completo(texto_original)
                    if USAR_FILTRO_GENERO:
                        genero = MAPA_GENERO_PERSONAGEM.get(personagem, "neutro")
                        traducao_final = aplicar_filtro_genero(traducao_final, genero)
                    registrar_problemas(texto_original, traducao_final, f"{os.path.basename(caminho_arquivo)}:{i + 1}")
                    linha_traduzida_formatada = f'{m_vazio_diag.group("indent")}{prefixo_vazio} "{traducao_final}"\n'
            
            elif m_com_narr and m_vazio_narr and m_com_narr.group('resto').strip() == m_vazio_narr.group('resto').strip():
                texto_original = m_com_narr.group('texto')
                traducao_final = processar_paragrafo_completo(texto_original)
                if USAR_FILTRO_GENERO:
                    traducao_final = aplicar_filtro_genero(traducao_final, "neutro")
                registrar_problemas(texto_original, traducao_final, f"{os.path.basename(caminho_arquivo)}:{i + 1}")
                linha_traduzida_formatada = f'{m_vazio_narr.group("indent")}"{traducao_final}"{m_vazio_narr.group("resto")}\n'
            
            if texto_original is not None and linha_traduzida_formatada is not None:
                novas_linhas.extend([linha_atual, linha_traduzida_formatada])
                traducoes_feitas += 1
                i += 2
                traduziu_bloco = True

        if not traduziu_bloco:
            novas_linhas.append(linha_atual)
            i += 1

    if traducoes_feitas > 0:
        caminho_backup = f"{caminho_arquivo}.bak"
        print(f"   -> {traducoes_feitas} linhas traduzidas. Criando backup: {os.path.basename(caminho_backup)}")
        try:
            shutil.copy(caminho_arquivo, caminho_backup)
            with open(caminho_arquivo, 'w', encoding='utf-8') as f:
                f.writelines(novas_linhas)
        except Exception as e:
            print(f"❌ Erro ao salvar o arquivo ou criar backup: {e}")
    else:
        print("   -> Nenhuma tradução necessária neste arquivo.")
    return traducoes_feitas

def modo_rpy(diretorio):
    """Função principal para o modo de tradução de arquivos Ren'Py."""
    print("\n--- MODO DE TRADUÇÃO REN'PY (.rpy) ---")
    if not os.path.isdir(diretorio):
        print(f"❌ ERRO: O diretório '{diretorio}' não existe.")
        sys.exit(1)

    global TRADUTOR
    TRADUTOR = configurar_tradutor("en", "pb", diretorio_base=diretorio) # Alterado de 'pb' para 'pt'

    arquivos_rpy = sorted([f for f in os.listdir(diretorio) if f.endswith(".rpy")])
    if not arquivos_rpy:
        print("⚠️  Nenhum arquivo .rpy encontrado no diretório.")
        return

    total_traducoes_geral = 0
    for nome_arquivo in arquivos_rpy:
        caminho_completo = os.path.join(diretorio, nome_arquivo)
        total_traducoes_geral += processar_arquivo_rpy(caminho_completo)

    print("\n" + "=" * 70)
    print("--- Processo Ren'Py Concluído ---")
    print(f"✅ Arquivos .rpy processados: {len(arquivos_rpy)}")
    print(f"✅ Total de linhas traduzidas: {total_traducoes_geral}")
    if RELATORIO_PROBLEMAS:
        print(f"⚠️  Encontrados {len(RELATORIO_PROBLEMAS)} possíveis problemas.")
        caminho_relatorio = os.path.join(diretorio, "relatorio_traducao.txt")
        with open(caminho_relatorio, 'w', encoding='utf-8') as f:
            for item in RELATORIO_PROBLEMAS:
                f.write(f"[{item['origem']}] {item['problemas']}\n")
                f.write(f"  ORIGINAL: {item['original']}\n")
                f.write(f"  TRADUZIDO: {item['traduzido']}\n\n")
        print(f"📄 Relatório salvo em: {caminho_relatorio}")
    else:
        print("✅ Nenhum problema detectado pelos filtros.")
    print("=" * 70)

# --- MODO DE TRADUÇÃO: WORD (.docx) ---

def modo_docx(caminho_arquivo):
    """Função principal para o modo de tradução de arquivos Word."""
    print("\n--- MODO DE TRADUÇÃO WORD (.docx) ---")
    if not os.path.isfile(caminho_arquivo):
        print(f"❌ ERRO: O arquivo '{caminho_arquivo}' não foi encontrado.")
        sys.exit(1)
    if not caminho_arquivo.lower().endswith('.docx'):
        print(f"❌ ERRO: O arquivo fornecido não é um .docx.")
        sys.exit(1)

    global TRADUTOR
    TRADUTOR = configurar_tradutor("en", "pb", diretorio_base=os.path.dirname(caminho_arquivo)) # Alterado de 'pb' para 'pt'

    try:
        print(f"📄 Lendo o arquivo: {os.path.basename(caminho_arquivo)}")
        documento_original = docx.Document(caminho_arquivo)
        documento_traduzido = docx.Document()
        total_paragrafos = len(documento_original.paragraphs)
        print(f"[*] Traduzindo {total_paragrafos} parágrafos...")

        for i, para in enumerate(documento_original.paragraphs):
            texto_para_traduzir = para.text
            if texto_para_traduzir.strip() == 'EMPTYSTRING':
                documento_traduzido.add_paragraph('EMPTYSTRING')
                continue

            if texto_para_traduzir.strip():
                traducao_final = processar_paragrafo_completo(texto_para_traduzir)
                if USAR_FILTRO_GENERO:
                    traducao_final = aplicar_filtro_genero(traducao_final, "neutro")
                registrar_problemas(texto_para_traduzir, traducao_final, f"{os.path.basename(caminho_arquivo)}:{i + 1}")
                documento_traduzido.add_paragraph(traducao_final)
            else:
                documento_traduzido.add_paragraph('')
            
            print(f"    -> Progresso: {i + 1}/{total_paragrafos}", end='\r')
        
        print("\n[*] Tradução concluída.")
        caminho_saida = os.path.join(os.path.dirname(caminho_arquivo), "untranslated pt-BR.docx")
        documento_traduzido.save(caminho_saida)

        print("\n" + "=" * 70)
        print("--- Processo DOCX Concluído ---")
        print(f"✅ Tradução salva em: {caminho_saida}")
        if RELATORIO_PROBLEMAS:
            caminho_relatorio = os.path.join(os.path.dirname(caminho_arquivo), "relatorio_traducao.txt")
            with open(caminho_relatorio, 'w', encoding='utf-8') as f:
                for item in RELATORIO_PROBLEMAS:
                    f.write(f"[{item['origem']}] {item['problemas']}\n")
                    f.write(f"  ORIGINAL: {item['original']}\n")
                    f.write(f"  TRADUZIDO: {item['traduzido']}\n\n")
            print(f"📄 Relatório salvo em: {caminho_relatorio}")
        print("=" * 70)

    except Exception as e:
        print(f"\n❌ OCORREU UM ERRO INESPERADO DURANTE O PROCESSAMENTO DOCX: {e}")
        sys.exit(1)

# --- INICIALIZAÇÃO DO SCRIPT ---

def iniciar_interface_grafica():
    janela = tk.Tk()
    janela.title("Super Tradutor Automático")
    janela.geometry("620x420")
    janela.configure(bg="#1f1f1f")

    estilo = ttk.Style()
    estilo.theme_use("clam")
    estilo.configure("TButton", font=("Segoe UI", 10), padding=6)
    estilo.configure("TLabel", font=("Segoe UI", 10), background="#1f1f1f", foreground="#f0f0f0")
    estilo.configure("TEntry", font=("Segoe UI", 10))

    caminho_var = tk.StringVar()
    modo_var = tk.StringVar(value="rpy")
    glossario_var = tk.StringVar()
    genero_var = tk.StringVar()
    padroes_var = tk.StringVar()
    filtro_var = tk.BooleanVar(value=True)

    def escolher_caminho():
        if modo_var.get() == "rpy":
            caminho = filedialog.askdirectory(title="Selecione a pasta do jogo")
        else:
            caminho = filedialog.askopenfilename(title="Selecione o arquivo .docx", filetypes=[("Word", "*.docx")])
        if caminho:
            caminho_var.set(caminho)

    def escolher_glossario():
        caminho = filedialog.askopenfilename(title="Selecione o glossário", filetypes=[("Texto", "*.txt")])
        if caminho:
            glossario_var.set(caminho)

    def escolher_genero():
        caminho = filedialog.askopenfilename(title="Selecione o mapa de gênero", filetypes=[("Texto", "*.txt")])
        if caminho:
            genero_var.set(caminho)

    def escolher_padroes():
        caminho = filedialog.askopenfilename(title="Selecione os padrões regex", filetypes=[("Texto", "*.txt")])
        if caminho:
            padroes_var.set(caminho)

    def executar():
        caminho = caminho_var.get().strip()
        if not caminho:
            messagebox.showerror("Erro", "Selecione um caminho válido.")
            return
        configurar_opcoes(
            glossario_var.get().strip(),
            genero_var.get().strip(),
            padroes_var.get().strip(),
            filtro_var.get()
        )
        try:
            if modo_var.get() == "rpy":
                modo_rpy(caminho)
            else:
                modo_docx(caminho)
            messagebox.showinfo("Concluído", "Processo finalizado. Veja o terminal para detalhes.")
        except SystemExit:
            pass
        except Exception as e:
            messagebox.showerror("Erro", f"Ocorreu um erro: {e}")

    frame = ttk.Frame(janela)
    frame.pack(fill="both", expand=True, padx=20, pady=20)

    ttk.Label(frame, text="Modo:").grid(row=0, column=0, sticky="w", pady=4)
    ttk.Radiobutton(frame, text="Ren'Py (.rpy)", variable=modo_var, value="rpy").grid(row=0, column=1, sticky="w")
    ttk.Radiobutton(frame, text="Word (.docx)", variable=modo_var, value="docx").grid(row=0, column=2, sticky="w")

    ttk.Label(frame, text="Caminho:").grid(row=1, column=0, sticky="w", pady=4)
    ttk.Entry(frame, textvariable=caminho_var, width=45).grid(row=1, column=1, columnspan=2, sticky="we")
    ttk.Button(frame, text="Selecionar", command=escolher_caminho).grid(row=1, column=3, padx=6)

    ttk.Label(frame, text="Glossário (opcional):").grid(row=2, column=0, sticky="w", pady=4)
    ttk.Entry(frame, textvariable=glossario_var, width=45).grid(row=2, column=1, columnspan=2, sticky="we")
    ttk.Button(frame, text="Selecionar", command=escolher_glossario).grid(row=2, column=3, padx=6)

    ttk.Label(frame, text="Mapa de gênero (opcional):").grid(row=3, column=0, sticky="w", pady=4)
    ttk.Entry(frame, textvariable=genero_var, width=45).grid(row=3, column=1, columnspan=2, sticky="we")
    ttk.Button(frame, text="Selecionar", command=escolher_genero).grid(row=3, column=3, padx=6)

    ttk.Label(frame, text="Padrões regex (opcional):").grid(row=4, column=0, sticky="w", pady=4)
    ttk.Entry(frame, textvariable=padroes_var, width=45).grid(row=4, column=1, columnspan=2, sticky="we")
    ttk.Button(frame, text="Selecionar", command=escolher_padroes).grid(row=4, column=3, padx=6)

    ttk.Checkbutton(frame, text="Aplicar filtro anti-feminino", variable=filtro_var).grid(row=5, column=0, columnspan=3, sticky="w", pady=6)

    ttk.Button(frame, text="Executar", command=executar).grid(row=6, column=0, columnspan=4, pady=20)

    for i in range(4):
        frame.columnconfigure(i, weight=1)

    janela.mainloop()

def configurar_opcoes(caminho_glossario, caminho_genero, caminho_padroes, usar_filtro):
    global GLOSSARIO_PROTEGIDO, MAPA_GENERO_PERSONAGEM, USAR_FILTRO_GENERO, RELATORIO_PROBLEMAS, PADROES_PROTEGIDOS
    GLOSSARIO_PROTEGIDO = carregar_glossario(caminho_glossario)
    MAPA_GENERO_PERSONAGEM = carregar_mapa_genero(caminho_genero)
    PADROES_PROTEGIDOS = PADROES_PROTEGIDOS_PADRAO + carregar_padroes_protegidos(caminho_padroes)
    USAR_FILTRO_GENERO = usar_filtro
    RELATORIO_PROBLEMAS = []

def main():
    parser = argparse.ArgumentParser(
        description="Super Tradutor Automático (v7.3) para arquivos .rpy e .docx.",
        formatter_class=argparse.RawTextHelpFormatter,
        epilog="""
Exemplos de uso:
  - Para traduzir arquivos Ren'Py em uma pasta:
    python seu_script.py --modo rpy ./caminho/para/pasta_do_jogo

  - Para traduzir um único arquivo .docx:
    python seu_script.py --modo docx ./caminho/para/meu_documento.docx
"""
    )
    parser.add_argument(
        'caminho',
        type=str,
        nargs='?',
        default=None,
        help="O caminho para o diretório (modo rpy) ou arquivo (modo docx)."
    )
    parser.add_argument(
        '--modo',
        type=str,
        choices=['rpy', 'docx', 'gui'],
        required=True,
        help="Define o modo de operação: 'rpy' para diretório Ren'Py ou 'docx' para um único arquivo Word."
    )
    parser.add_argument(
        '--glossario',
        type=str,
        default=None,
        help="Caminho para um glossário de termos que não devem ser traduzidos."
    )
    parser.add_argument(
        '--genero',
        type=str,
        default=None,
        help="Caminho para um arquivo de mapa de gênero por personagem."
    )
    parser.add_argument(
        '--padroes',
        type=str,
        default=None,
        help="Caminho para um arquivo com padrões regex adicionais a serem protegidos."
    )
    parser.add_argument(
        '--sem-filtro-genero',
        action='store_true',
        help="Desativa o filtro anti-feminino."
    )

    args = parser.parse_args()

    if args.modo == 'gui':
        iniciar_interface_grafica()
        return

    if not args.caminho:
        print("❌ ERRO: Caminho obrigatório para os modos rpy/docx.")
        sys.exit(1)

    configurar_opcoes(args.glossario, args.genero, args.padroes, not args.sem_filtro_genero)

    if args.modo == 'rpy':
        modo_rpy(args.caminho)
    elif args.modo == 'docx':
        modo_docx(args.caminho)

if __name__ == "__main__":
    main()
